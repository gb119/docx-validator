#!/usr/bin/env python
"""
Demo script to show enhanced logging functionality.

This script demonstrates the new debug logging that captures:
1. Full prompts sent to the LLM
2. Full responses from the LLM  
3. HTTP response metadata
4. Token usage information
"""

import logging
import sys
import tempfile
from pathlib import Path

# Set up logging to show debug messages
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    stream=sys.stdout
)

# Import after setting up logging
from docx_tex_validator import DocxValidator, ValidationSpec

def create_sample_docx():
    """Create a simple sample DOCX file for testing."""
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed, using mock file")
        return None
    
    doc = Document()
    doc.core_properties.title = "Sample Document"
    doc.core_properties.author = "Test Author"
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is a sample document for testing.")
    
    # Save to temp file
    temp_file = tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False)
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name


def main():
    print("=" * 80)
    print("Enhanced Logging Demo - docx-tex-validator")
    print("=" * 80)
    print()
    print("This demo shows the new comprehensive debug logging that captures:")
    print("  - Full LLM prompts (complete conversation)")
    print("  - Full LLM responses")
    print("  - Token usage information")
    print("  - Response metadata")
    print()
    print("Note: This is a mock demo (no actual API calls)")
    print("=" * 80)
    print()
    
    # Create a mock validator (won't make actual API calls)
    print("Creating validator with gpt-4o (new default model with 128k context)...")
    validator = DocxValidator(
        model_name="gpt-4o",  # New default with larger context window
        api_key="mock-key-for-demo"
    )
    print(f"✓ Validator created with model: {validator.backend.model_name}")
    print()
    
    # Show that the new default model is gpt-4o
    print("Default Model Information:")
    print(f"  Model: {validator.backend.model_name}")
    print(f"  Backend: {validator.backend.name}")
    print(f"  Context Window: 128k tokens (much larger than gpt-4o-mini)")
    print()
    
    print("To see actual debug logging in action:")
    print("1. Set OPENAI_API_KEY environment variable")
    print("2. Run validation with logging.DEBUG level enabled")
    print("3. Check the logs for:")
    print("   - 'LLM REQUEST' sections showing full prompts")
    print("   - 'Response data' showing full LLM responses")
    print("   - 'Token usage' showing token consumption")
    print("   - 'Response metadata' showing HTTP response details")
    print()
    print("Example command:")
    print("  python -c 'import logging; logging.basicConfig(level=logging.DEBUG); ...")
    print()
    print("=" * 80)


if __name__ == "__main__":
    main()
