"""
Parser for .docx files (Microsoft Word documents).
"""

import zipfile
from typing import Any, Dict

from docx import Document

from .base import BaseParser


class DocxParser(BaseParser):
    """Parser for extracting structure and metadata from .docx files.

    Examples:
        >>> parser = DocxParser()
        >>> structure = parser.parse("document.docx")
        >>> print(structure['metadata']['title'])
    """

    def supports_extension(self, extension: str) -> bool:
        """Check if this parser supports the given file extension.

        Args:
            extension (str):
                File extension to check.

        Returns:
            (bool):
                True if extension is '.docx', False otherwise.
        """
        return extension.lower() == ".docx"

    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse a .docx file and extract its structure and content information.

        Args:
            file_path (str):
                Path to the .docx file.

        Returns:
            (Dict[str, Any]):
                Dictionary containing document structure information.

        Raises:
            FileNotFoundError:
                If the file doesn't exist.
            ValueError:
                If the file is not a valid .docx file.
        """
        self.validate_file(file_path)

        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"Failed to parse .docx file: {e}") from e

        # Extract document structure
        structure = {
            "file_path": str(file_path),
            "document_type": "docx",
            "paragraphs": [],
            "tables": [],
            "sections": [],
            "styles": set(),
            "has_header": False,
            "has_footer": False,
            "metadata": {},
            "xml_content": None,
        }

        # Extract paragraph information
        for para in doc.paragraphs:
            para_info = {
                "text": para.text,
                "style": para.style.name if para.style else None,
                "alignment": str(para.alignment) if para.alignment else None,
                "runs_count": len(para.runs),
            }
            structure["paragraphs"].append(para_info)
            if para.style:
                structure["styles"].add(para.style.name)

        # Extract table information
        for table in doc.tables:
            table_info = {
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": [],
            }
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_info["cells"].append(row_data)
            structure["tables"].append(table_info)

        # Extract section information
        for section in doc.sections:
            section_info = {
                "page_width": section.page_width.inches if section.page_width else None,
                "page_height": section.page_height.inches if section.page_height else None,
                "orientation": str(section.orientation)
                if hasattr(section, "orientation")
                else None,
            }
            structure["sections"].append(section_info)

        # Check for headers and footers
        if doc.sections:
            first_section = doc.sections[0]
            structure["has_header"] = bool(first_section.header.paragraphs)
            structure["has_footer"] = bool(first_section.footer.paragraphs)

        # Extract metadata
        core_props = doc.core_properties
        structure["metadata"] = {
            "author": core_props.author or "",
            "title": core_props.title or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else None,
            "modified": str(core_props.modified) if core_props.modified else None,
        }

        # Convert styles set to list for JSON serialization
        structure["styles"] = list(structure["styles"])

        # Extract raw XML content from the DOCX file for advanced validation
        # This allows the LLM to inspect cross-references, field codes, captions, etc.
        try:
            with zipfile.ZipFile(file_path, "r") as docx_zip:
                if "word/document.xml" in docx_zip.namelist():
                    xml_bytes = docx_zip.read("word/document.xml")
                    # Use errors='replace' to handle malformed UTF-8 gracefully
                    structure["xml_content"] = xml_bytes.decode("utf-8", errors="replace")
        except (zipfile.BadZipFile, UnicodeDecodeError, OSError):
            # If XML extraction fails (bad ZIP, encoding issues, or I/O error),
            # continue without it - the basic structure information is still available
            pass

        return structure

    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        """Legacy method for backward compatibility.

        This method is deprecated. Use parse() instead.

        Args:
            file_path (str):
                Path to the .docx file.

        Returns:
            (Dict[str, Any]):
                Dictionary containing document structure information.
        """
        parser = DocxParser()
        return parser.parse(file_path)
