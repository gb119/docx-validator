"""
Module for parsing .docx files and extracting document structure information.
"""

import zipfile
from pathlib import Path
from typing import Any, Dict

from docx import Document


class DocxParser:
    """Parser for extracting structure and metadata from .docx files."""

    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        """
        Parse a .docx file and extract its structure and content information.

        Args:
            file_path: Path to the .docx file

        Returns:
            Dictionary containing document structure information

        Raises:
            FileNotFoundError: If the file doesn't exist
            ValueError: If the file is not a valid .docx file
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not path.suffix.lower() == ".docx":
            raise ValueError(f"File must be a .docx file, got: {path.suffix}")

        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"Failed to parse .docx file: {e}") from e

        # Extract document structure
        structure = {
            "file_path": str(file_path),
            "paragraphs": [],
            "tables": [],
            "sections": [],
            "styles": set(),
            "has_header": False,
            "has_footer": False,
            "metadata": {},
            "xml_content": None,
        }

        # Extract paragraph information
        for para in doc.paragraphs:
            para_info = {
                "text": para.text,
                "style": para.style.name if para.style else None,
                "alignment": str(para.alignment) if para.alignment else None,
                "runs_count": len(para.runs),
            }
            structure["paragraphs"].append(para_info)
            if para.style:
                structure["styles"].add(para.style.name)

        # Extract table information
        for table in doc.tables:
            table_info = {
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": [],
            }
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_info["cells"].append(row_data)
            structure["tables"].append(table_info)

        # Extract section information
        for section in doc.sections:
            section_info = {
                "page_width": section.page_width.inches if section.page_width else None,
                "page_height": section.page_height.inches if section.page_height else None,
                "orientation": str(section.orientation)
                if hasattr(section, "orientation")
                else None,
            }
            structure["sections"].append(section_info)

        # Check for headers and footers
        if doc.sections:
            first_section = doc.sections[0]
            structure["has_header"] = bool(first_section.header.paragraphs)
            structure["has_footer"] = bool(first_section.footer.paragraphs)

        # Extract metadata
        core_props = doc.core_properties
        structure["metadata"] = {
            "author": core_props.author or "",
            "title": core_props.title or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else None,
            "modified": str(core_props.modified) if core_props.modified else None,
        }

        # Convert styles set to list for JSON serialization
        structure["styles"] = list(structure["styles"])

        # Extract raw XML content from the DOCX file for advanced validation
        # This allows the LLM to inspect cross-references, field codes, captions, etc.
        try:
            with zipfile.ZipFile(file_path, "r") as docx_zip:
                if "word/document.xml" in docx_zip.namelist():
                    xml_bytes = docx_zip.read("word/document.xml")
                    # Use errors='replace' to handle malformed UTF-8 gracefully
                    structure["xml_content"] = xml_bytes.decode("utf-8", errors="replace")
        except (zipfile.BadZipFile, UnicodeDecodeError, OSError):
            # If XML extraction fails (bad ZIP, encoding issues, or I/O error),
            # continue without it - the basic structure information is still available
            pass

        return structure
